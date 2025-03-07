import os
import time
from tkinter import filedialog

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.oxml import qn
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt


def input_L3word(L3_filepath):
    # 创建一个Word文档
    L3_doc = Document(L3_filepath)
    img_files = ['img_FBD.png','img_GLQFBD.png','img_IPS.png','img_LDSM.png','img_RZSJ.png','img_WAF.png','img_WLSJ.png']
    # 图片文件列表
    image_files = ['CDU-2.png', 'word_L3JKZY.png', 'word_L3JKGJ.png',
                   'word_L3-512-1电源.png', 'word_L3-512-1总有功功率.png',
                   'word_L3-512-2电源.png', 'word_L3-512-2总有功功率.png',
                   'word_L3-512-3电源.png', 'word_L3-512-3总有功功率.png',
                   'word_L3-512-4电源.png', 'word_L3-512-4总有功功率.png',
                   '核心机房-温湿度-北面微模块-1.png', '核心机房-温湿度-南面微模块-1.png',
                   '温湿度-冷冻室西北.png', '温湿度-电气室西北.png',
                   "L3平台机柜监控图.png", "L3第一套立方体监控视图.png", "L3第二套立方体监控视图.png",
                   "L3第三套立方体监控视图.png", "L3第四套立方体监控视图.png",
                   'word_L3YPT.png', 'word_LDSM.png', 'word_RZSJ.png', 'word_WLSJ.png',
                   'word_FDB.png', 'word_GLQFDB.png','word_IPS.png', 'word_WAF.png']

    # 单元格位置列表
    cell_positions = [(6, 1), (13, 1), (13, 1), (14, 1), (14, 1), (15, 1), (15, 1), (16, 1), (16, 1), (17, 1), (17, 1),
                      (18, 1), (18, 1), (19, 1), (20, 1), (21, 1), (22, 1), (23, 1), (24, 1), (25, 1),
                      (26, 1), (27, 1), (28, 1), (29, 1), (30, 1), (31, 1), (32, 1),(33, 1)]

    # 遍历图片和对应位置，批量删除图片
    for image_file, cell_position in zip(image_files, cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        table = L3_doc.tables[1]  # 获取第一个表格
        cell = table.cell(cell_position[0], cell_position[1])  # 获取指定单元格
        cell.text = ''

    # 遍历图片和对应位置，批量插入图片
    for image_file, cell_position in zip(image_files, cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        table = L3_doc.tables[1]  # 获取第一个表格
        cell = table.cell(cell_position[0], cell_position[1])  # 获取指定单元格
        # 将图片插入到单元格
        cell.paragraphs[0].add_run().add_picture(image_file, width=Inches(3.8), height=Inches(2.2))

        # 获取文档中的第一个表格
    L3date_time = L3_doc.tables[0]
    # 获取指定单元格并写入当前日期
    L3write_time = L3date_time.cell(1, 4)
    L3write_time.text = time.strftime("%Y.%m.%d")
    # 设置字体为仿宋，大小为四号（14 磅）
    paragraph = L3write_time.paragraphs[0]
    for run in paragraph.runs:
        run.font.size = Pt(12)  # 小四字体对应 12 磅
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.bold = True  # 设置字体加粗
        # 确保设置居中对齐生效
    cell_align = L3date_time.cell(1, 4)
    for paragraph in cell_align.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 生成文档保存的文件名
    L3date_name = time.strftime("%Y%m%d")
    L3res_doc = 'L3项目运维日报-' + L3date_name
    # 弹出保存文件对话框
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", ".docx")],
                                             initialfile=L3res_doc)
    # 根据用户选择保存文档或取消操作
    if save_path:
        L3_doc.save(save_path)
        print(f"Document saved to {save_path}")
    else:
        print("Save operation cancelled.")
    # 去除图片文件列表中的重复项
    L3all_picture = list(set(image_files + img_files))

    for L3rm_picture in L3all_picture:
        os.remove(L3rm_picture)

def input_L2word(L2_filepath):
    # 创建一个Word文档
    L2_doc = Document(L2_filepath)

    # 图片文件列表
    L2image_files = ["L2-液冷4.png", 'word_L2JKZY.png', 'word_L2JKGJ.png', "word_L2-512-1电源.png",
                     "word_L2-512-1总有功功率.png", "word_L2-512-2电源.png", "word_L2-512-2总有功功率.png",
                     "L2平台机柜监控图.png",
                     "L2第一套立方体监控视图.png", "L2第二套立方体监控视图.png", "L2-512-1热通道温湿度.png",
                     "L2-512-2热通道温湿度.png",
                     "L2-温湿度-电池上.png", "L2-冷冻室温湿度.png"]

    # 单元格位置列表
    L2cell_positions = [(7, 1), (11, 1), (11, 1), (12, 1), (12, 1), (13, 1), (13, 1), (14, 1), (15, 1), (16, 1), (17, 1),
                        (18, 1), (19, 1), (20, 1)]

    # 遍历图片和对应位置，批量删除图片
    for L2image_file, L2cell_position in zip(L2image_files, L2cell_positions):
        L2table = L2_doc.tables[1]  # 获取第一个表格
        L2cell = L2table.cell(L2cell_position[0], L2cell_position[1])  # 获取指定单元格
        L2cell.text = ''

    # 遍历图片和对应位置，批量插入图片
    for L2image_file, L2cell_position in zip(L2image_files, L2cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        L2table = L2_doc.tables[1]  # 获取第一个表格
        L2cell = L2table.cell(L2cell_position[0], L2cell_position[1])  # 获取指定单元格
        # 将图片插入到单元格
        L2cell.paragraphs[0].add_run().add_picture(L2image_file, width=Inches(3.8), height=Inches(2.2))

        # 获取文档中的第一个表格
    L2date_time = L2_doc.tables[0]
    # 获取指定单元格并写入当前日期
    L2write_time = L2date_time.cell(1, 3)
    L2write_time.text = time.strftime("%Y.%m.%d")
    # 设置字体为仿宋，大小为四号（14 磅）
    paragraph = L2write_time.paragraphs[0]
    for run in paragraph.runs:
        run.font.size = Pt(12)  # 小四字体对应 12 磅
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.bold = True  # 设置字体加粗
        # 确保设置居中对齐生效
    cell_align = L2date_time.cell(1, 3)
    for paragraph in cell_align.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 生成文档保存的文件名
    L2date_name = time.strftime("%Y%m%d")
    L2res_doc = 'L2项目运维日报-' + L2date_name
    # 弹出保存文件对话框
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", ".docx")],
                                             initialfile=L2res_doc)
    # 根据用户选择保存文档或取消操作
    if save_path:
        L2_doc.save(save_path)
        print(f"Document saved to {save_path}")
    else:
        print("Save operation cancelled.")
    # 去除图片文件列表中的重复项
    L2all_picture = list(set(L2image_files))

    for L2rm_picture in L2all_picture:
        os.remove(L2rm_picture)


def input_L5word(L5_filepath):
    L5_doc = Document(L5_filepath)

    # 图片文件列表
    L5image_files = ["L5-CDU-1.png", "L5-虚拟机.png", "L5-虚机存储.png", "word_L5JKZY.png", "word_L5KGJ.png",
                     "word_L5-512-智能电表.png",
                     "word_L5-512-总有功功率.png", "L5-温湿度-512-RACK1.png", "L5-温湿度-512-RACK2.png",
                     "L5-温湿度-512-RACK3.png",
                     "L5-温湿度-512-RACK4.png", "L5-温湿度-动力配套室东.png", "L5-温湿度-水泵房.png", "L5-门禁监控.png"]

    # 单元格位置列表
    L5cell_positions = [(7, 1), (10, 1), (10, 1), (14, 1), (14, 1), (15, 1), (15, 1), (16, 1), (17, 1), (18, 1), (19, 1), (20, 1),
                        (21, 1), (22, 1)]
    # 遍历图片和对应位置，批量删除图片
    for L5image_file, L5cell_position in zip(L5image_files, L5cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        L5table = L5_doc.tables[1]  # 获取第一个表格
        L5cell = L5table.cell(L5cell_position[0], L5cell_position[1])  # 获取指定单元格
        L5cell.text = ''
    # 遍历图片和对应位置，批量插入图片
    for L5image_file, L5cell_position in zip(L5image_files, L5cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        L5table = L5_doc.tables[1]  # 获取第一个表格
        L5cell = L5table.cell(L5cell_position[0], L5cell_position[1])  # 获取指定单元格
        # 将图片插入到单元格
        L5cell.paragraphs[0].add_run().add_picture(L5image_file, width=Inches(3.8), height=Inches(2.2))

    L5date_time = L5_doc.tables[0]
    # 获取指定单元格并写入当前日期
    L5write_time = L5date_time.cell(1, 3)
    L5write_time.text = time.strftime("%Y.%m.%d")
    # 设置字体为仿宋，大小为四号（14 磅）
    paragraph = L5write_time.paragraphs[0]
    for run in paragraph.runs:
        run.font.size = Pt(12)  # 小四字体对应 12 磅
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.bold = True  # 设置字体加粗
        # 确保设置居中对齐生效
    cell_align = L5date_time.cell(1, 3)
    for paragraph in cell_align.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 生成文档保存的文件名
    L5date_name = time.strftime("%Y%m%d")
    L5res_doc = 'L5项目运维日报-' + L5date_name
    # 弹出保存文件对话框
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", ".docx")],
                                             initialfile=L5res_doc)
    # 根据用户选择保存文档或取消操作
    if save_path:
        L5_doc.save(save_path)
        print(f"Document saved to {save_path}")
    else:
        print("Save operation cancelled.")
    # 去除图片文件列表中的重复项
    L5all_picture = list(set(L5image_files))

    for L5rm_picture in L5all_picture:
        os.remove(L5rm_picture)


def input_S2word(S2_filepath):
    S2_doc = Document(S2_filepath)

    # 图片文件列表
    S2image_files = ["S2-虚拟机.png", "S2-虚机存储.png", "word_S2JKZY.png", "S2-CDU(150KW)监控.png", "S2-冷水机组监控.png",
                     "S2-通信电源监控1.png", "S2-通信电源监控2.png", "S2-通信电源监控3.png", "S2-通信电源监控4.png",
                     "S2-通信电源监控5.png","S2-通信电源监控6.png", "S2-通信电源监控7.png", "S2-通信电源监控8.png",
                     "S2-门禁记录.png", "S2-实时告警.png",'S2-监控前视图.png','S2-监控后视图.png']

    # 单元格位置列表
    S2cell_positions = [(5, 1), (5, 1), (7, 1), (8, 1), (9, 1), (10, 1), (11, 1), (12, 1), (13, 1), (14, 1), (15, 1), (16, 1),
                        (17, 1), (18, 1), (19, 1), (20, 1), (21, 1)]
    # 遍历图片和对应位置，批量插入图片
    for S2image_file, S2cell_position in zip(S2image_files, S2cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        S2table = S2_doc.tables[1]  # 获取第一个表格
        S2cell = S2table.cell(S2cell_position[0], S2cell_position[1])  # 获取指定单元格
        S2cell.text = ''
    # 遍历图片和对应位置，批量插入图片
    for S2image_file, S2cell_position in zip(S2image_files, S2cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        S2table = S2_doc.tables[1]  # 获取第一个表格
        S2cell = S2table.cell(S2cell_position[0], S2cell_position[1])  # 获取指定单元格
        # 将图片插入到单元格
        S2cell.paragraphs[0].add_run().add_picture(S2image_file, width=Inches(3.8), height=Inches(2.2))

    # 获取文档中的第一个表格
    S2date_time = S2_doc.tables[0]
    # 获取指定单元格并写入当前日期
    S2write_time = S2date_time.cell(1, 3)
    S2write_time.text = time.strftime("%Y.%m.%d")
    # 设置字体为仿宋，大小为四号（14 磅）
    paragraph = S2write_time.paragraphs[0]
    for run in paragraph.runs:
        run.font.size = Pt(12)  # 小四字体对应 12 磅
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.bold = True  # 设置字体加粗
        # 确保设置居中对齐生效
    cell_align = S2date_time.cell(1, 3)
    for paragraph in cell_align.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 生成文档保存的文件名
    S2date_name = time.strftime("%Y%m%d")
    S2res_doc = 'S2项目运维日报-' + S2date_name
    # 弹出保存文件对话框
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", ".docx")],
                                             initialfile=S2res_doc)
    # 根据用户选择保存文档或取消操作
    if save_path:
        S2_doc.save(save_path)
        print(f"Document saved to {save_path}")
    else:
        print("Save operation cancelled.")
    # 去除图片文件列表中的重复项
    S2all_picture = list(set(S2image_files))

    for S2rm_picture in S2all_picture:
        os.remove(S2rm_picture)


def input_HPC64word(HPC64_filepath):
    HPC64_doc = Document(HPC64_filepath)

    # 图片文件列表
    HPC64image_files = ["HPC64-液冷机柜.png", "HPC64-CDU(35KW)监控.png", "HPC64-水冷空调监控.png", "HPC64-通信电源监控1.png",
                        "HPC64-通信电源监控2.png", "HPC64-实时告警.png"]

    # 单元格位置列表
    HPC64cell_positions = [(8, 1), (9, 1), (10, 1), (11, 1), (12, 1), (13, 1), (14, 1)]
    # 遍历图片和对应位置，批量插入图片
    for HPC64image_file, HPC64cell_position in zip(HPC64image_files, HPC64cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        HPC64table = HPC64_doc.tables[1]  # 获取第一个表格
        HPC64cell = HPC64table.cell(HPC64cell_position[0], HPC64cell_position[1])  # 获取指定单元格
        HPC64cell.text = ''
    # 遍历图片和对应位置，批量插入图片
    for HPC64image_file, HPC64cell_position in zip(HPC64image_files, HPC64cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        HPC64table = HPC64_doc.tables[1]  # 获取第一个表格
        HPC64cell = HPC64table.cell(HPC64cell_position[0], HPC64cell_position[1])  # 获取指定单元格
        # 将图片插入到单元格
        HPC64cell.paragraphs[0].add_run().add_picture(HPC64image_file, width=Inches(3.8), height=Inches(2.2))

        # 获取文档中的第一个表格
    HPC64date_time = HPC64_doc.tables[0]
    # 获取指定单元格并写入当前日期
    HPC64write_time = HPC64date_time.cell(1, 3)
    HPC64write_time.text = time.strftime("%Y.%m.%d")
    # 设置字体为仿宋，大小为四号（14 磅）
    paragraph = HPC64write_time.paragraphs[0]
    for run in paragraph.runs:
        run.font.size = Pt(12)  # 小四字体对应 12 磅
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.bold = True  # 设置字体加粗
        # 确保设置居中对齐生效
    cell_align = HPC64date_time.cell(1, 3)
    for paragraph in cell_align.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 生成文档保存的文件名
    HPC64date_name = time.strftime("%Y%m%d")
    HPC64res_doc = '64节点一体机项目运维日报-' + HPC64date_name
    # 弹出保存文件对话框
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", ".docx")],
                                             initialfile=HPC64res_doc)
    # 根据用户选择保存文档或取消操作
    if save_path:
        HPC64_doc.save(save_path)
        print(f"Document saved to {save_path}")
    else:
        print("Save operation cancelled.")
    # 去除图片文件列表中的重复项
    HPC64all_picture = list(set(HPC64image_files))

    for HPC64rm_picture in HPC64all_picture:
        os.remove(HPC64rm_picture)


def input_L6word(L6_filepath):
    L6_doc = Document(L6_filepath)

    # 图片文件列表
    L6image_files = ["L6-冷水机组实时数据.png", "L6-CDU实时数据.png", "L6-二次侧机柜流速.png", "word_L6GJ.png",
                     "word_L6-512智能电表-核心系统.png", "word_L6-512-总有功功率.png", "L6-冷水机组.png", "L6-CDU.png",
                     "L6ZZJD-虚拟机.png", "L6ZZJD-虚机存储.png"]

    # 单元格位置列表
    L6cell_positions = [(6, 1), (7, 1), (8, 1), (13, 1), (14, 1), (14, 1), (15, 1), (16, 1), (17, 1), (17, 1)]
    # 遍历图片和对应位置，批量删除图片
    for L6image_file, L6cell_position in zip(L6image_files, L6cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        L6table = L6_doc.tables[1]  # 获取第一个表格
        L6cell = L6table.cell(L6cell_position[0], L6cell_position[1])  # 获取指定单元格
        L6cell.text = ''
    # 遍历图片和对应位置，批量插入图片
    for L6image_file, L6cell_position in zip(L6image_files, L6cell_positions):
        # 插入图片到Word文档的指定表格位置
        # 假设表格已经存在，这里我们只插入图片
        L6table = L6_doc.tables[1]  # 获取第一个表格
        L6cell = L6table.cell(L6cell_position[0], L6cell_position[1])  # 获取指定单元格
        # 将图片插入到单元格
        L6cell.paragraphs[0].add_run().add_picture(L6image_file, width=Inches(3.8), height=Inches(2.2))

        # 获取文档中的第一个表格
    L6date_time = L6_doc.tables[0]
    # 获取指定单元格并写入当前日期
    L6write_time = L6date_time.cell(1, 3)
    L6write_time.text = time.strftime("%Y.%m.%d")
    # 设置字体为仿宋，大小为四号（14 磅）
    paragraph = L6write_time.paragraphs[0]
    for run in paragraph.runs:
        run.font.size = Pt(12)  # 小四字体对应 12 磅
        run.font.name = '仿宋'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.bold = True  # 设置字体加粗
        # 确保设置居中对齐生效
    cell_align = L6date_time.cell(1, 3)
    for paragraph in cell_align.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 生成文档保存的文件名
    L6date_name = time.strftime("%Y%m%d")
    L6res_doc = '激光产业园总装基地运维日报-' + L6date_name
    # 弹出保存文件对话框
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", ".docx")],
                                             initialfile=L6res_doc)
    # 根据用户选择保存文档或取消操作
    if save_path:
        L6_doc.save(save_path)
        print(f"Document saved to {save_path}")
    else:
        print("Save operation cancelled.")
    # 去除图片文件列表中的重复项
    L6all_picture = list(set(L6image_files))

    for L6rm_picture in L6all_picture:
        os.remove(L6rm_picture)


if __name__ == "__main__":
    print('')
